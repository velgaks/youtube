import youtube_dl
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from docx import Document

# List of YouTube video URLs
video_urls = ['https://www.youtube.com/watch?v=VPZD_aij8H0', # has English subs provided 
              'https://www.youtube.com/watch?v=A7XvTkzsvY8'] # has auto generated subs in Ukr

# Create a new Word document
document = Document()
formatter = TextFormatter()

for video_url in video_urls:
    with youtube_dl.YoutubeDL() as ydl:
        video_info = ydl.extract_info(video_url, download=False)
        video_title = video_info.get('title')
        video_id = video_url.split("=")[-1]
        transcript = YouTubeTranscriptApi.get_transcript(video_id, 
                                                     languages=['uk', 'en'])        
        t = formatter.format_transcript(transcript)
        t = t.replace('\n', ' ') # remove line breaks
        document.add_heading(video_title, level=1)
        document.add_paragraph(t)
    # Add a new page before saving the document
    document.add_page_break()

# Save the document
document.save('transcripts_and_subtitles.docx')